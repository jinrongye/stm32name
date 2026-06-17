"""
夜间智能避障系统 — 实验报告生成器
运行: python 生成实验报告.py
输出: 实验报告.docx
"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

BASE = os.path.dirname(os.path.abspath(__file__))
USER = os.path.join(BASE, "43", "1", "user")

def read_code(filename):
    """读取代码文件，返回内容字符串"""
    path = os.path.join(USER, filename)
    if os.path.exists(path):
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    return f"/* 文件 {filename} 未找到 */"

# ===== 读取关键代码 =====
code_main     = read_code("main.c")
code_hc_sr04_c = read_code("hc_sr04.c")
code_hc_sr04_h = read_code("hc_sr04.h")
code_oled_c   = read_code("oled.c")

def extract_function(code, func_name):
    """从代码中提取指定函数（从函数定义到上一个 } 结束）"""
    lines = code.split("\n")
    start = -1
    for i, line in enumerate(lines):
        if func_name in line and ("void " in line or "static " in line or "float " in line or "uint" in line):
            # 找到函数定义行
            if "(" in line and ")" in line and ("{" in line or lines[i+1].strip() == "{"):
                start = i
                break
    if start < 0:
        return f"/* {func_name} 未找到 */"
    # 从 start 开始，找到匹配的 }
    brace = 0
    started = False
    result = []
    for i in range(start, len(lines)):
        result.append(lines[i])
        if "{" in lines[i]:
            brace += lines[i].count("{")
            started = True
        if "}" in lines[i]:
            brace -= lines[i].count("}")
        if started and brace == 0:
            break
    return "\n".join(result)

# ===== 提取关键函数片段 =====
fn_measure       = extract_function(code_hc_sr04_c, "HC_SR04_Measure")
fn_init_sr04     = extract_function(code_hc_sr04_c, "HC_SR04_Init")
fn_night_display = extract_function(code_main, "Display_NightMode")
fn_day_display   = extract_function(code_main, "Display_DayMode")
fn_adc_init      = extract_function(code_main, "ADC1_Init")
fn_adc_read      = extract_function(code_main, "ADC1_Read")
fn_main_loop     = extract_function(code_main, "main")

# ===== 创建 Word 文档 =====
doc = Document()

# --- 全局样式设置 ---
style = doc.styles["Normal"]
font = style.font
font.name = "宋体"
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(4)

# 标题样式
for i in range(1, 4):
    h_style = doc.styles[f"Heading {i}"]
    h_font = h_style.font
    h_font.name = "黑体"
    h_font.color.rgb = RGBColor(0, 0, 0)
    if i == 1:
        h_font.size = Pt(18)
    elif i == 2:
        h_font.size = Pt(15)
    else:
        h_font.size = Pt(13)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_para(text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "宋体"
    run.font.size = Pt(12)
    return p

def add_code_block(code_text, title=None):
    """添加代码块（灰色背景框）"""
    if title:
        p = doc.add_paragraph()
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(11)
        r.font.name = "宋体"
    # 用表格模拟代码块
    lines = code_text.strip().split("\n")
    # 限制代码块长度（取前60行或全部）
    display_lines = lines[:55]
    if len(lines) > 55:
        display_lines.append(f"/* ... 省略 {len(lines)-55} 行 ... */")

    for line in display_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.1
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    # 代码块后空一行
    doc.add_paragraph()

# ===== 封面 =====
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("夜间智能避障系统\n—— 实验报告")
r.bold = True
r.font.name = "黑体"
r.font.size = Pt(26)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_text = (
    "课程：嵌入式系统设计\n"
    "主控芯片：STM32F103C8T6\n"
    "开发环境：Keil MDK-ARM\n"
    "日期：2026年6月17日"
)
r = info.add_run(info_text)
r.font.name = "宋体"
r.font.size = Pt(14)

doc.add_page_break()

# ===== 目录占位 =====
add_heading("目  录", level=1)
add_para("（在 Word 中按 Ctrl+A → F9 刷新目录域）")
doc.add_page_break()

# =====================================================================
#  第一章：实验目的
# =====================================================================
add_heading("一、实验目的", level=1)

add_para(
    "本实验旨在设计并实现一个基于 STM32F103C8T6 微控制器的夜间智能避障系统，"
    "主要实验目的如下：", indent=True
)

purposes = [
    "掌握 STM32F103C8T6 的 GPIO、ADC、定时器等外设的配置与使用方法；",
    "学习 HC-SR04 超声波测距模块的工作原理及驱动编程；",
    "掌握 SSD1306 OLED 显示屏的 I2C 通信协议及软件模拟 I2C 的实现；",
    "理解光敏传感器的工作原理，实现环境光照强度的模拟信号采集；",
    "综合运用多传感器（光敏 + 超声波）实现智能判断逻辑——"
    "环境变暗时自动开启超声波测距，环境明亮时自动待机；",
    "培养嵌入式系统软硬件联调能力，掌握 Keil MDK 集成开发环境的使用。"
]
for p in purposes:
    add_para("• " + p)

# =====================================================================
#  第二章：实验运用场景
# =====================================================================
add_heading("二、实验运用场景", level=1)

add_para(
    "本系统模拟了夜间环境下的障碍物检测与避障预警功能，具有广泛的实际应用价值：", indent=True
)

scenes = [
    ("智能家居", "夜间自动检测走廊或房间内的障碍物，辅助夜灯照明控制，避免老人或儿童在黑暗中绊倒。"),
    ("智能机器人", "作为机器人夜间巡航的避障子系统，当环境光线不足时自动激活超声波传感器检测前方障碍物。"),
    ("安防监控", "夜间检测到不明物体靠近时发出警告，可用于家庭或仓库的低成本安防预警。"),
    ("智能停车场", "夜间检测车位是否有车辆停靠，或引导车辆安全泊入。"),
    ("盲人辅助设备", "穿戴式设备在夜间自动开启超声波探测，提示前方障碍物距离。"),
]
for title, desc in scenes:
    add_para(f"【{title}】{desc}", indent=True)

# =====================================================================
#  第三章：实验原理
# =====================================================================
add_heading("三、实验原理", level=1)

# 3.1 系统总体架构
add_heading("3.1 系统总体架构", level=2)
add_para(
    "本系统以 STM32F103C8T6 为核心控制器，通过光敏传感器采集环境光照强度（ADC），"
    "判断当前处于日间还是夜间。当检测到黑暗环境时，自动启动 HC-SR04 超声波传感器测距，"
    "并将距离信息显示在 SSD1306 OLED 屏幕上。同时，板载 LED（PC13）在夜间模式下点亮作为指示灯。",
    indent=True
)

# 3.2 光敏传感器原理
add_heading("3.2 光敏传感器原理", level=2)
add_para(
    "光敏传感器模块采用光敏电阻（LDR）配合 LM393 电压比较器。光敏电阻的阻值随光照强度变化："
    "光照越强，电阻越小；光照越弱，电阻越大。通过分压电路将电阻变化转换为电压变化，"
    "输出至 STM32 的 ADC1 通道0（PA0）进行 12 位模数转换。"
    "ADC 采样值范围为 0~4095，对应 0~3.3V 电压。"
    "代码中设定电压阈值 1.5V 作为日间/夜间的分界线。",
    indent=True
)

add_code_block(fn_adc_init, "【ADC1 初始化代码 — ADC1_Init()】")
add_code_block(fn_adc_read, "【ADC 读取代码 — ADC1_Read()】")

# 3.3 超声波测距原理
add_heading("3.3 HC-SR04 超声波测距原理", level=2)
add_para(
    "HC-SR04 模块采用超声波回波测距法。工作流程如下：",
    indent=True
)
steps_ultrasonic = [
    "MCU 向 TRIG 引脚发送一个 ≥10μs 的高电平脉冲（本代码使用约 15μs）；",
    "HC-SR04 模块接收到触发信号后，自动发射 8 个 40kHz 的超声波脉冲；",
    "超声波遇到障碍物后反射回来，被模块的接收探头捕获；",
    "模块将 ECHO 引脚拉高，高电平持续时间等于超声波从发射到接收的往返时间；",
    "MCU 使用 TIM3 定时器（配置为 1MHz 自由计数器）测量 ECHO 高电平的精确持续时间；",
    "根据声速公式换算距离：声速 ≈ 340 m/s = 0.034 cm/μs，"
    "距离(cm) = (脉宽(μs) × 0.034) / 2 ≈ 脉宽(μs) / 58。"
]
for i, s in enumerate(steps_ultrasonic, 1):
    add_para(f"{i}. {s}")

add_para(
    "HC-SR04 主要参数：工作电压 5V，测量范围 2cm ~ 400cm，"
    "测量角度约 15°，精度约 3mm。",
    indent=True
)

add_code_block(fn_init_sr04, "【HC-SR04 初始化代码 — HC_SR04_Init()】")
add_code_block(fn_measure, "【超声波测距核心代码 — HC_SR04_Measure()】")

# 时序图（ASCII art）
add_heading("3.4 超声波工作时序", level=2)
timing = (
    "TRIG:   ___|‾‾‾‾‾‾‾‾‾‾‾‾‾‾‾|___________________________\n"
    "           ←  ≥10μs (代码约15μs)  →\n\n"
    "        模块内部发射 8 个 40kHz 超声波脉冲\n"
    "           ←      约 200μs      →\n\n"
    "ECHO:   _______|‾‾‾‾‾‾‾‾‾‾‾‾‾‾‾‾‾‾‾‾‾|______________\n"
    "               ←  高电平脉宽 ∝ 距离  →\n"
    "               距离 = 脉宽(μs) / 58 (cm)"
)
add_code_block(timing, "【HC-SR04 时序图】")

# 3.5 OLED 显示原理
add_heading("3.5 SSD1306 OLED 显示原理", level=2)
add_para(
    "SSD1306 是一款 128×64 像素的单色 OLED 驱动芯片，本实验使用 I2C 通信接口。"
    "由于 STM32 的硬件 I2C 在某些情况下不稳定，本设计采用 GPIO 软件模拟 I2C 时序："
    "PB6 模拟 SCL（时钟线），PB7 模拟 SDA（数据线），均配置为开漏输出。"
    "I2C 地址为 0x78（8位写地址，对应 7 位地址 0x3C）。",
    indent=True
)
add_para(
    "显示使用 8×16 点阵英文字体（size=16），每行显示 16 个英文字符。"
    "OLED 屏幕垂直方向 64 像素，分三行均匀排列（y=0, y=24, y=48），行间距 8 像素。",
    indent=True
)

# =====================================================================
#  第四章：实验步骤与代码
# =====================================================================
add_heading("四、实验步骤及对应代码", level=1)

add_heading("4.1 硬件连接", level=2)
add_para(
    "按照下表完成各模块与 STM32F103C8T6 的引脚连接：", indent=True
)

# 硬件连接表
table = doc.add_table(rows=8, cols=4, style="Light Grid Accent 1")
headers = ["STM32 引脚", "连接至", "信号类型", "功能说明"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)

connections = [
    ("PA0", "光敏模块 DO", "模拟输入 (ADC1_CH0)", "读取环境光照电压"),
    ("PB0", "HC-SR04 TRIG", "推挽输出", "发送触发脉冲"),
    ("PB1", "HC-SR04 ECHO", "浮空输入 (TIM3)", "接收回波信号"),
    ("PB6", "OLED SCL", "开漏输出", "I2C 时钟线"),
    ("PB7", "OLED SDA", "开漏输出", "I2C 数据线"),
    ("PC13", "板载 LED", "推挽输出", "夜间指示灯"),
    ("5V", "HC-SR04 VCC", "电源", "⚠ 必须 5V，不可 3.3V"),
]
for i, (pin, to, sig, desc) in enumerate(connections, 1):
    table.rows[i].cells[0].text = pin
    table.rows[i].cells[1].text = to
    table.rows[i].cells[2].text = sig
    table.rows[i].cells[3].text = desc

doc.add_paragraph()

add_heading("4.2 软件初始化", level=2)
add_para(
    "系统上电后，依次执行以下初始化流程：", indent=True
)
init_steps = [
    "SysTick_Init() — 配置系统滴答定时器，产生 1ms 中断作为全局时基；",
    "OLED_Init() — 初始化 SSD1306 OLED 显示屏（I2C 软件模拟），发送初始化命令序列；",
    "LED_Init() — 配置 PC13 为推挽输出，默认高电平（LED 灭）；",
    "ADC1_Init() — 配置 ADC1 通道0（PA0），单次转换模式，软件触发；",
    "HC_SR04_Init() — 配置 PB0(TRIG) 推挽输出、PB1(ECHO) 浮空输入、TIM3 1MHz 自由计数器；",
    "显示启动画面 — OLED 显示 \"Night Avoid System Ready\" 约 1 秒后进入主循环。"
]
for i, s in enumerate(init_steps, 1):
    add_para(f"{i}. {s}")

add_heading("4.3 主循环逻辑", level=2)
add_para(
    "主循环每 100ms 执行一次，核心判断逻辑如下：", indent=True
)

logic = (
    "while (1) {\n"
    "    adc_value = ADC1_Read();                        // 读取光敏传感器 ADC 值\n"
    "    voltage = adc_value * 3.3 / 4095;               // 换算为电压值\n"
    "\n"
    "    if (voltage >= 1.5) {                           // 遮挡→黑暗→夜间模式\n"
    "        GPIO_ResetBits(PC13);                       // LED 亮\n"
    "        distance = HC_SR04_Measure();               // 超声波测距\n"
    "        Display_NightMode(distance);                // OLED 显示夜间界面\n"
    "    } else {                                        // 光照→明亮→日间模式\n"
    "        GPIO_SetBits(PC13);                         // LED 灭\n"
    "        Display_DayMode();                          // OLED 显示日间界面\n"
    "    }\n"
    "    Delay_ms(100);\n"
    "}"
)
add_code_block(logic, "【主循环伪代码】")

add_code_block(fn_main_loop, "【main.c 主函数完整代码】")

add_heading("4.4 OLED 显示函数", level=2)
add_code_block(fn_night_display, "【夜间模式显示 — Display_NightMode()】")
add_code_block(fn_day_display, "【日间模式显示 — Display_DayMode()】")

# =====================================================================
#  第五章：实验结果
# =====================================================================
add_heading("五、实验结果", level=1)

add_heading("5.1 日间模式测试", level=2)
add_para(
    "在正常室内光照条件下，光敏传感器输出电压 ≥ 1.5V（ADC 值约 1860 以上），"
    "系统判定为日间模式。此时：",
    indent=True
)
day_results = [
    "板载 LED（PC13）熄灭；",
    "HC-SR04 超声波模块处于待机状态，不发送触发脉冲；",
    "OLED 显示 \"DAY MODE\" 和 \"Ultrasonic OFF\"，字体清晰，布局合理。"
]
for r in day_results:
    add_para("• " + r)

add_heading("5.2 夜间模式测试", level=2)
add_para(
    "用手遮挡光敏传感器后，输出电压降至 1.5V 以下，系统自动切换为夜间模式。此时：",
    indent=True
)
night_results = [
    "板载 LED（PC13）亮起，指示夜间工作状态；",
    "HC-SR04 自动触发测距，TIM3 1MHz 计数器精确测量 ECHO 脉宽；",
    "OLED 显示 \"NIGHT MODE\"、实时距离值（格式：XX.X cm）、安全状态（Safe/DANGER!）；",
    "在不同距离放置障碍物（如书本），OLED 显示的距离值能正确响应变化；",
    "当障碍物距离 < 30cm 时，OLED 显示 \"DANGER!\" 危险警告。"
]
for r in night_results:
    add_para("• " + r)

add_heading("5.3 OLED 界面展示", level=2)

displays = [
    ("日间模式",
     "┌──────────────────────┐\n"
     "│                      │\n"
     "│     DAY MODE         │  ← 16px大字，居中\n"
     "│                      │\n"
     "│  Ultrasonic OFF      │  ← 16px大字\n"
     "│                      │\n"
     "└──────────────────────┘"),
    ("夜间模式（正常）",
     "┌──────────────────────┐\n"
     "│    NIGHT MODE        │  ← 16px大字，居中\n"
     "│  Dist: 125.3 cm      │  ← 实时距离\n"
     "│        Safe          │  ← 安全提示\n"
     "└──────────────────────┘"),
    ("夜间模式（危险）",
     "┌──────────────────────┐\n"
     "│    NIGHT MODE        │\n"
     "│  Dist:  18.5 cm      │  ← < 30cm\n"
     "│     DANGER!          │  ← 危险警告\n"
     "└──────────────────────┘"),
]
for title, ascii in displays:
    add_para(f"【{title}】", bold=True)
    add_code_block(ascii)

# =====================================================================
#  第六章：思考与结论
# =====================================================================
add_heading("六、思考与结论", level=1)

add_heading("6.1 实验过程中遇到的问题及解决方法", level=2)

problems = [
    ("OLED 屏幕不亮",
     "I2C 软件模拟的延时过短（原约 2μs），部分 SSD1306 模块的 I2C 时序较慢，"
     "导致通信失败。将延时增加至约 10μs 后显示正常。此外，增加上电等待时间确保 SSD1306 内部电路稳定。"),
    ("光敏传感器检测模式反转",
     "原代码假设\"遮挡=低压=夜间\"，但实际光敏模块的 LM393 比较器输出特性相反——"
     "遮挡时光敏电阻增大，DO 输出电压升高。将判断条件从 voltage < 1.5 改为 voltage >= 1.5 后恢复正常。"),
    ("HC-SR04 始终返回 No Sig",
     "最初采用 TIM3 输入捕获 + 中断方式测量 ECHO 脉宽，但中断可能因 NVIC 配置或输入滤波设置问题未能正常触发。"
     "重构为 GPIO 轮询 + TIM3 自由计数器 + SysTick 超时的简化方案后，测距功能正常。"
     "此问题也提醒我们：超声波模块必须使用 5V 供电，STM32 的 3.3V 无法驱动。"),
    ("OLED 显示布局拥挤",
     "初始方案使用 6×8 小字体、行间距不足，导致信息显示拥挤。"
     "改用 8×16 大字体，三行数据均匀分布在 64 像素高度中（y=0, 24, 48），视觉效果显著改善。"),
]
for title, desc in problems:
    add_para(f"【问题】{title}", bold=True)
    add_para(f"【解决】{desc}", indent=True)
    doc.add_paragraph()

add_heading("6.2 实验收获", level=2)
gains = [
    "深入理解了 STM32 标准外设库（SPL）的 GPIO、ADC、TIM、NVIC 等外设的配置方法；",
    "掌握了 HC-SR04 超声波传感器的工作原理：电平触发、回波脉宽与距离的线性关系；",
    "学会了 SSD1306 OLED 的 I2C 通信协议及软件模拟 I2C 的实现技巧；",
    "体会了软硬件联调的重要性——代码逻辑正确不代表系统能正常工作，硬件供电、"
    "接线、时序匹配同样关键；",
    "认识到嵌入式系统中\"简化即可靠\"的设计哲学——从复杂的中断输入捕获方案简化为"
    "GPIO 轮询方案后，可靠性反而大幅提升；",
    "培养了工程调试思维：遇到问题首先排查硬件（供电、接线），再排查软件（时序、逻辑）。"
]
for g in gains:
    add_para("• " + g)

add_heading("6.3 改进方向", level=2)
improvements = [
    "增加蜂鸣器报警功能 — 当距离 < 30cm 时发出声音警告；",
    "接入 SIM900A 模块 — 检测到异常时发送短信报警；",
    "使用 NRF24L01 无线模块 — 将测距数据无线传输至远端主机；",
    "增加数据记录功能 — 使用 SD 卡或 Flash 记录历史测距数据；",
    "引入 FreeRTOS 实时操作系统 — 将传感器读取、显示刷新、报警处理分为独立任务。"
]
for imp in improvements:
    add_para("• " + imp)

add_heading("6.4 结论", level=2)
add_para(
    "本实验成功设计并实现了一个基于 STM32F103C8T6 的夜间智能避障系统。"
    "系统能够根据环境光照强度自动切换工作模式——日间超声波待机省电，"
    "夜间自动开启超声波测距并在 OLED 上显示距离信息，障碍物距离小于 30cm 时触发危险警告。"
    "实验验证了多传感器融合、嵌入式外设驱动开发、实时数据显示等关键技术，"
    "达到了预期的实验目标。",
    indent=True
)

# ===== 附录 =====
doc.add_page_break()
add_heading("附录：硬件模块清单", level=1)
table2 = doc.add_table(rows=7, cols=3, style="Light Grid Accent 1")
for i, h in enumerate(["序号", "模块名称", "型号/规格"]):
    table2.rows[0].cells[i].text = h
    for r in table2.rows[0].cells[i].paragraphs[0].runs:
        r.bold = True
hw_list = [
    ("1", "主控板", "STM32F103C8T6"),
    ("2", "OLED 显示屏", "SSD1306 128×64 I2C"),
    ("3", "光敏感应模块", "光敏电阻 + LM393"),
    ("4", "超声波测距模块", "HC-SR04"),
    ("5", "USB-TTL 下载器", "ST-Link / CH340G"),
    ("6", "杜邦线及面包板", "公母头杜邦线若干"),
]
for i, (n, name, model) in enumerate(hw_list, 1):
    table2.rows[i].cells[0].text = n
    table2.rows[i].cells[1].text = name
    table2.rows[i].cells[2].text = model

doc.add_paragraph()
add_heading("附录：引脚连接速查表", level=2)
table3 = doc.add_table(rows=7, cols=2, style="Light Grid Accent 1")
table3.rows[0].cells[0].text = "STM32 引脚"
table3.rows[0].cells[1].text = "连接目标"
for r in table3.rows[0].cells:
    for run in r.paragraphs[0].runs:
        run.bold = True
pins = [
    ("PA0", "光敏模块 DO"),
    ("PB0", "HC-SR04 TRIG"),
    ("PB1", "HC-SR04 ECHO"),
    ("PB6", "OLED SCL"),
    ("PB7", "OLED SDA"),
    ("PC13", "板载 LED"),
]
for i, (pin, target) in enumerate(pins, 1):
    table3.rows[i].cells[0].text = pin
    table3.rows[i].cells[1].text = target

# ===== 保存 =====
output_path = os.path.join(BASE, "实验报告.docx")
doc.save(output_path)
print(f"[OK] 实验报告已生成: {output_path}")
